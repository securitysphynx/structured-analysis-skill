"""
Extract full text content from all .docx files in the gemini-analysis directory.

Usage:
    uv run helper_scripts/extract_docx_text.py

Requires: python-docx (installed via uv pip install python-docx)
"""

import docx
import os


def extract_docx_text(directory: str) -> None:
    """Extract and print text from all .docx files in the given directory."""
    files = sorted(os.listdir(directory))

    for fname in files:
        if not fname.endswith(".docx"):
            continue
        filepath = os.path.join(directory, fname)
        doc = docx.Document(filepath)

        print("=" * 80)
        print(f"FILE: {fname}")
        print("=" * 80)

        for para in doc.paragraphs:
            print(para.text)

        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                print(" | ".join(cells))

        print()
        print()


if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    gemini_dir = os.path.join(base_dir, "docs", "background", "gemini-analysis")
    extract_docx_text(gemini_dir)
